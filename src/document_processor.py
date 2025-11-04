import os
import fitz  # PyMuPDF
from docx import Document
from typing import Union, Tuple
import re

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        
    def process_document(self, file_path: str) -> Tuple[str, str]:
        """
        Process the uploaded document and extract text with better formatting.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Tuple[str, str]: (extracted text, file extension)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format. Supported formats are: {', '.join(self.supported_formats)}")
            
        if file_ext == '.pdf':
            return self._process_pdf(file_path), file_ext
        elif file_ext == '.docx':
            return self._process_docx(file_path), file_ext
        else:  # .txt
            return self._process_txt(file_path), file_ext

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace while preserving sentence structure
        text = re.sub(r'\n+', '\n', text)  # Multiple newlines to single
        text = re.sub(r' +', ' ', text)    # Multiple spaces to single
        
        # Fix common OCR issues
        text = text.replace('�', ' ')  # Remove replacement characters
        
        # Ensure proper sentence separation
        text = re.sub(r'\.([A-Z])', r'. \1', text)  # Add space after period if missing
        text = re.sub(r'\. +', '. ', text)  # Normalize spaces after periods
        
        return text.strip()
            
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with better formatting preservation and error handling."""
        doc = None
        try:
            doc = fitz.open(file_path)
            text_blocks = []
            
            for page_num, page in enumerate(doc):
                try:
                    # Try to get text with layout preservation
                    text = page.get_text("text")
                    if text and text.strip():
                        text_blocks.append(text)
                    else:
                        print(f"Warning: Page {page_num + 1} appears to be an image or has no extractable text")
                        # Try alternative text extraction method
                        try:
                            text = page.get_text("dict")
                            if text and 'blocks' in text:
                                page_text = ""
                                for block in text['blocks']:
                                    if 'lines' in block:
                                        for line in block['lines']:
                                            if 'spans' in line:
                                                for span in line['spans']:
                                                    if 'text' in span:
                                                        page_text += span['text'] + " "
                                if page_text.strip():
                                    text_blocks.append(page_text)
                        except Exception as e:
                            print(f"    Alternative extraction failed for page {page_num + 1}: {e}")
                            
                except Exception as e:
                    print(f"Error processing page {page_num + 1}: {e}")
                    continue
            
            if not text_blocks:
                print("Warning: No text could be extracted from any pages")
                return "Error: Could not extract text from PDF"
            
            full_text = '\n'.join(text_blocks)
            cleaned_text = self._clean_text(full_text)
            
            print(f"Successfully extracted {len(cleaned_text)} characters from {len(text_blocks)} pages")
            return cleaned_text
            
        except Exception as e:
            print(f"Critical error processing PDF: {e}")
            return "Error: Could not process PDF file"
        finally:
            if doc:
                try:
                    doc.close()
                except:
                    pass
        
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with enhanced structure preservation."""
        doc = Document(file_path)
        text_blocks = []
        
        print(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")
        
        # Extract paragraphs with better sentence preservation
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                # Ensure proper sentence separation
                if not para_text.endswith('.') and not para_text.endswith('!') and not para_text.endswith('?'):
                    # Only add period if it looks like a sentence (not a header/title)
                    if len(para_text) > 20 and not para_text.isupper():
                        para_text += '.'
                
                text_blocks.append(para_text)
        
        # Extract text from tables with better formatting
        for table_idx, table in enumerate(doc.tables):
            print(f"Processing table {table_idx + 1}")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Clean up cell text
                        cell_text = re.sub(r'\s+', ' ', cell_text)
                        row_text.append(cell_text)
                
                if row_text:
                    # Join cells with proper punctuation
                    combined_row = '. '.join(row_text)
                    if not combined_row.endswith('.'):
                        combined_row += '.'
                    text_blocks.append(combined_row)
        
        # Join with proper paragraph separation
        full_text = '\n\n'.join(text_blocks)
        cleaned_text = self._clean_text(full_text)
        
        print(f"Extracted {len(cleaned_text)} characters from DOCX")
        return cleaned_text
        
    def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT file with encoding detection."""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    return self._clean_text(text)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            text = file.read()
            return self._clean_text(text)

    def get_document_stats(self, text: str) -> dict:
        """Get basic statistics about the processed document"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        words = text.split()
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len([p for p in text.split('\n') if p.strip()]),
        }
            
    def save_highlighted_document(self, file_path: str, highlighted_text: str, output_path: str) -> None:
        """
        Save the highlighted document.
        
        Args:
            file_path (str): Original file path
            highlighted_text (str): Text with highlighting markers
            output_path (str): Path to save the highlighted document
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            self._save_highlighted_pdf(file_path, highlighted_text, output_path)
        elif file_ext == '.docx':
            self._save_highlighted_docx(file_path, highlighted_text, output_path)
        else:  # .txt
            self._save_highlighted_txt(highlighted_text, output_path)
            
    def _save_highlighted_pdf(self, file_path: str, highlighted_text: str, output_path: str) -> None:
        """Save highlighted PDF document."""
        doc = fitz.open(file_path)
        # Implementation for PDF highlighting will be handled by DocumentHighlighter
        doc.save(output_path)
        doc.close()
        
    def _save_highlighted_docx(self, file_path: str, highlighted_text: str, output_path: str) -> None:
        """Save highlighted DOCX document."""
        doc = Document(file_path)
        # Implementation for DOCX highlighting will be handled by DocumentHighlighter
        doc.save(output_path)
        
    def _save_highlighted_txt(self, highlighted_text: str, output_path: str) -> None:
        """Save highlighted TXT document."""
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(highlighted_text)